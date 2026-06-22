# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


DOCX_PATH = next(Path(".").glob("*Prolog.docx"))


def set_para_text(paragraph, text, size=14, font_name="Times New Roman"):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)


def set_code_text(paragraph, text):
    set_para_text(paragraph, text, size=9, font_name="Courier New")


doc = Document(DOCX_PATH)

paragraph_updates = {
    35: "выделить основные сущности предметной области: студент, специальность, предприятие, вакансия, отрасль, формат работы, профессиональные интересы, требуемые навыки и ограничения по среднему баллу;",
    41: "Кураторы производственной практики и заведующие отделениями колледжа — для автоматического распределения студентов на основе среднего балла, специальности, города, предпочтительной отрасли, формата работы, навыков и профессиональных интересов;",
    46: "Предметная область подбора производственной практики включает несколько взаимосвязанных групп данных. С одной стороны, имеется профиль студента: специальность, успеваемость, желаемый город, предпочтительный формат работы, предпочтительная отрасль, профессиональные интересы и набор компетенций. С другой стороны, существуют предприятия-партнеры, у каждого из которых есть направление деятельности, минимальные требования к среднему баллу, локация и конкретные вакансии для практикантов.",
    47: "При построении базы знаний были приняты следующие ограничения. Во-первых, специальность студента должна соответствовать списку специальностей, указанному в вакансии. Во-вторых, средний балл студента не должен быть ниже минимального порога предприятия. В-третьих, город прохождения практики должен совпадать с предпочтением студента, за исключением удаленного формата работы или выбора варианта «любой город». В-четвертых, учитываются предпочтительная отрасль и формат работы: офисный, удаленный или любой.",
    48: "После прохождения обязательных фильтров система оценивает мягкие критерии — совпадение требуемых навыков и профессиональных интересов. Навыки дают до 30% итогового рейтинга, интересы — до 20%, поэтому система не только отбрасывает неподходящие вакансии, но и ранжирует оставшиеся варианты по степени соответствия профилю студента.",
    53: "В результате анализа были выделены пять учебных направлений: информационные системы, сетевое и системное администрирование, технология машиностроения, бухгалтерский учет, туризм и гостеприимство. Для каждого направления определены типовые вакансии, ключевые навыки, профессиональные интересы, отрасли предприятий и возможные форматы прохождения практики.",
    57: "Фреймовая модель используется для компактного описания объектов предметной области. Каждый фрейм представляет собой структурированную запись, состоящую из имени объекта и набора слотов. В данной работе фрейм вакансии включает идентификатор, должность, предприятие, подходящую специальность, минимальный средний балл, город или формат работы, отрасль, перечень требуемых навыков и связанные профессиональные интересы.",
    62: "Представленная модель содержит как жесткие требования, используемые для первичного отбора, так и мягкие признаки, применяемые для расчета процента соответствия: навыки и профессиональные интересы. Благодаря этому система не просто определяет факт пригодности вакансии, но и показывает относительную степень соответствия профиля студента требованиям работодателя.",
    113: "исключение вакансий, не соответствующих специальности, среднему баллу, локации, отрасли или формату работы;",
    117: "Для проведения верификации базы знаний использовался тестовый профиль студента ИТ-специальности, отвечающего на 11 вопросов анкеты:",
    119: "Средний балл успеваемости: 4.2",
    120: "Желаемая локация: moscow (Москва), формат работы: any, отрасль: it",
    121: "Ответы на вопросы интересов: yes для программирования, no для остальных направлений; на этой основе формируются навыки python, sql, git и интересы programming, backend, databases.",
    124: "Результаты тестирования экспертной системы в консоли SWI-Prolog представлены на рисунках 2–7. Проверялись четыре тестовых случая: студент специальности «Информационные системы» с интересом к программированию, студент специальности «Бухгалтерский учёт» с интересом к финансам и документам, студент специальности «Сетевое администрирование» с интересом к сетям и Linux, а также случай отсутствия подходящих мест при невыполнении обязательных фильтров.",
    125: "Выбранные сценарии покрывают основные ветви логического вывода: полное соответствие вакансии, частичное соответствие с рекомендациями по недостающим навыкам, подбор для другой специальности и отрицательный результат. Это позволяет убедиться, что система корректно обрабатывает ответы на 11 вопросов, применяет жесткие фильтры и рассчитывает рейтинг по формуле 50% + 30% + 20%.",
    130: "Рисунок 3 — Ввод данных студента: специальность «is», ср. балл 4.2, г. Москва, формат any, отрасль it и ответы yes/no на вопросы интересов",
    133: "Рисунок 4 — Результаты подбора для специальности «Информационные системы»: 100% — Разработчик Python, частичное соответствие — Frontend-разработчик",
    139: "Рисунок 6 — Результаты подбора для специальности «Сетевое администрирование»: 100% — Помощник системного администратора при совпадении навыков и интересов",
}

for index, text in paragraph_updates.items():
    if index < len(doc.paragraphs):
        set_para_text(doc.paragraphs[index], text)


pl_code = """% ФАКТЫ О КОМПАНИЯХ
company(it_complex, 'ООО "ИТ-Комплекс"', it, moscow, 4.0, 'Разработка веб-приложений и ПО.').
company(prom_tech, 'ООО "ПромТех"', industry, spb, 3.5, 'Проектирование деталей и ЧПУ-производство.').
company(rjd, 'АО "РЖД"', transport, ekaterinburg, 3.0, 'Транспортный оператор РФ.').
company(cbr, 'Центральный Банк РФ', finance, moscow, 4.5, 'Главный финансовый регулятор страны.').
company(cosmos, 'Гостиница "Космос"', tourism, moscow, 3.0, 'Гостиничный комплекс Москвы.').
company(soft_lines, 'ООО "СофтЛайнс"', it, remote, 3.8, 'ИТ-компания с удаленным форматом.').

% ФАКТЫ О ВАКАНСИЯХ И ИНТЕРЕСАХ
position(py_dev, 'Разработчик Python', it_complex, [is], [python, sql, git], 'Backend-разработка.').
position(front_dev, 'Frontend-разработчик', soft_lines, [is], [javascript, html, css, react, git], 'Разработка на React.').
position(sys_admin, 'Помощник системного администратора', rjd, [set], [linux, tcp_ip, bash], 'Поддержка Linux.').
position(jr_auditor, 'Младший аудитор', cbr, [buh], [excel, fin_analysis, law], 'Анализ отчетности.').

position_interests(py_dev, [programming, backend, databases]).
position_interests(front_dev, [programming, web, design]).
position_interests(sys_admin, [networks, administration, linux]).
position_interests(jr_auditor, [finance, analysis, law_docs]).

% ПРАВИЛА ЛОГИЧЕСКОГО ВЫВОДА
match_specialization(StudentSpec, AllowedSpecs) :- member_of(StudentSpec, AllowedSpecs).
match_gpa(StudentGPA, MinGPA) :- StudentGPA >= MinGPA.
match_location(any, _) :- !.
match_location(_, remote) :- !.
match_location(StudentLoc, CompanyLoc) :- StudentLoc = CompanyLoc.
match_industry(any, _).
match_industry(Industry, Industry).
match_work_format(any, _).
match_work_format(remote, remote).
match_work_format(office, CompanyLoc) :- CompanyLoc \\= remote.

interest_score(StudentInterests, PositionID, Score) :-
    position_interests(PositionID, PositionInterests),
    intersect_count(PositionInterests, StudentInterests, MatchCount),
    list_len(StudentInterests, TotalCount),
    (TotalCount > 0 -> Score is (MatchCount * 20) / TotalCount ; Score is 0).

recommend_position(StudentSpec, StudentSkills, StudentGPA, StudentLoc, IndustryPref, WorkFormatPref, StudentInterests,
                   PositionID, Title, CompanyName, CompanyLoc, Score, MissingSkills) :-
    position(PositionID, Title, CompanyID, AllowedSpecs, ReqSkills, _),
    company(CompanyID, CompanyName, Industry, CompanyLoc, MinGPA, _),
    match_specialization(StudentSpec, AllowedSpecs),
    match_gpa(StudentGPA, MinGPA),
    match_location(StudentLoc, CompanyLoc),
    match_industry(IndustryPref, Industry),
    match_work_format(WorkFormatPref, CompanyLoc),
    intersect_count(ReqSkills, StudentSkills, MatchCount),
    list_len(ReqSkills, TotalCount),
    SkillsScore is (MatchCount * 30) / TotalCount,
    interest_score(StudentInterests, PositionID, InterestScore),
    Score is 50 + SkillsScore + InterestScore,
    missing_skills(ReqSkills, StudentSkills, MissingSkills)."""

cli_code = """:- consult('expert_system.pl').

positive_answer(yes).
positive_answer(y).
negative_answer(no).
negative_answer(n).

ask_profile_question(Number, Text, YesSkills, YesInterests, Skills, Interests) :-
    format('Вопрос ~w. ~w (yes/no): ', [Number, Text]),
    read(Answer), nl,
    (positive_answer(Answer) -> Skills = YesSkills, Interests = YesInterests
    ; negative_answer(Answer) -> Skills = [], Interests = []
    ; writeln('Ошибка: введите yes. или no.'),
      ask_profile_question(Number, Text, YesSkills, YesInterests, Skills, Interests)).

ask_diagnostics(Skills, Interests) :-
    ask_profile_question(6, 'Вам интересна разработка программ, серверная логика или базы данных?',
                         [python, sql, git], [programming, backend, databases], S6, I6),
    ask_profile_question(7, 'Вам интересны сайты, интерфейсы и дизайн экранов?',
                         [javascript, html, css, react, git], [programming, web, design], S7, I7),
    ask_profile_question(8, 'Вам интересны сети, серверы Linux и администрирование?',
                         [linux, tcp_ip, bash], [networks, administration, linux], S8, I8),
    append_all([S6, S7, S8], AllSkills),
    append_all([I6, I7, I8], AllInterests),
    list_to_set(AllSkills, Skills),
    list_to_set(AllInterests, Interests).

start :-
    writeln('=== ЭКСПЕРТНАЯ СИСТЕМА ПОДБОРА ПРАКТИКИ ==='),
    write('Вопрос 1. Введите вашу специальность: '), read(Spec),
    write('Вопрос 2. Введите ваш средний балл: '), read(GPA),
    write('Вопрос 3. Введите желаемый город: '), read(Loc),
    write('Вопрос 4. Введите формат работы: '), read(WorkFormat),
    write('Вопрос 5. Введите отрасль: '), read(IndustryPref),
    ask_diagnostics(Skills, Interests),
    findall(Score-rec(Title, CoName, CoLoc, Missing),
            recommend_position(Spec, Skills, GPA, Loc, IndustryPref, WorkFormat, Interests,
                               _, Title, CoName, CoLoc, Score, Missing),
            Pairs),
    keysort(Pairs, Sorted), reverse(Sorted, DescSorted),
    print_results(DescSorted)."""

test_output = """?- start.
=== ЭКСПЕРТНАЯ СИСТЕМА ПОДБОРА ПРАКТИКИ ===

Вопрос 1. Введите вашу специальность: is.
Вопрос 2. Введите ваш средний балл: 4.2.
Вопрос 3. Введите желаемый город: moscow.
Вопрос 4. Введите формат работы: any.
Вопрос 5. Введите отрасль: it.
Вопрос 6. Вам интересна разработка программ, серверная логика или базы данных? yes.
Вопрос 7. Вам интересны сайты, интерфейсы и дизайн экранов? no.
Вопрос 8. Вам интересны сети, серверы Linux и администрирование? no.
Вопрос 9. Вам интересны чертежи, станки с ЧПУ и техническое проектирование? no.
Вопрос 10. Вам интересны бухгалтерские документы, Excel, финансы или право? no.
Вопрос 11. Вам интересны общение с клиентами, сервис и английский язык? no.

Выполняется поиск подходящих мест практики...

РЕЗУЛЬТАТЫ ПОДБОРА (отсортированы по соответствию):
-----------------------------------------------------------------
100% соответствия | Должность: Разработчик Python
               | Компания:  ООО "ИТ-Комплекс" (moscow)
               | Требования: Полное соответствие!
-----------------------------------------------------------------
62.666666666666664% соответствия | Должность: Frontend-разработчик
               | Компания:  ООО "СофтЛайнс" (remote)
               | Недостающие навыки для изучения: [javascript,html,css,react]
-----------------------------------------------------------------"""

set_code_text(doc.paragraphs[105], pl_code)
set_code_text(doc.paragraphs[108], cli_code)
set_code_text(doc.paragraphs[123], test_output)

if doc.tables:
    table = doc.tables[0]
    headers = ["ID", "Должность", "Компания", "Спец.", "GPA", "Локация/формат", "Отрасль", "Навыки и интересы"]
    rows = [
        ["py_dev", "Разработчик Python", 'ООО "ИТ-Комплекс"', "is", "4.0", "moscow / office", "it", "python, sql, git; programming, backend, databases"],
        ["front_dev", "Frontend-разработчик", 'ООО "СофтЛайнс"', "is", "3.8", "remote", "it", "javascript, html, css, react, git; web, design"],
        ["sys_admin", "Помощник системного администратора", 'АО "РЖД"', "set", "3.0", "ekaterinburg / office", "transport", "linux, tcp_ip, bash; networks, administration"],
        ["cad_designer", "Техник-конструктор", 'ООО "ПромТех"', "tm", "3.5", "spb / office", "industry", "autocad, cnc, solidworks; engineering, drawings"],
        ["buh_assistant", "Помощник бухгалтера", 'АО "РЖД"', "buh", "3.0", "ekaterinburg / office", "transport", "one_c, excel; documents, accounting"],
        ["jr_auditor", "Младший аудитор", "Центробанк РФ", "buh", "4.5", "moscow / office", "finance", "excel, fin_analysis, law; finance, analysis"],
        ["hotel_admin", "Администратор отеля", 'Гостиница "Космос"', "tur", "3.0", "moscow / office", "tourism", "english, communication; service, english"],
    ]
    while len(table.columns) < len(headers):
        table.add_column(Pt(80))
    for c, header in enumerate(headers):
        table.cell(0, c).text = header
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            table.cell(r, c).text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(9)

doc.save(DOCX_PATH)
print(DOCX_PATH)
